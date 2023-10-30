import docx
import numpy as np
import pandas as pd
import scipy.stats as stats
import json

from src.models.recommender import (
    Recommender_multinomial_logit,
    Recommender_most_popular,
    Recommender_random,
)
from src.evaluation.process_results import read_data_for_models
from src.data.load_dataset import load_dataset
from src.utils import create_path
from src.paths import EVALUATION_RESULTS_PATH, OUTPUT_TABLE_PATH
from src.config import OUTPUT_TABLES_MODEL_NAMES


def round_str(f, pos=2):
    """Round a float to a string with a given number of decimal places."""
    return str(np.round(f, pos))

def significance_symbol(p, symbol="*", n_tests=129):
    """Return a string with a significance symbol."""
    if p * n_tests > 0.1:
        return ""
    elif p * n_tests > 0.05:
        return symbol
    elif p * n_tests > 0.01:
        return 2 * symbol
    else:
        return 3 * symbol

def print_table(table):
    """Print a table."""
    print()
    print(pd.DataFrame(table).to_string(header=False, index=False))
    print()

def generate_and_print_bias_table(
    title, header, models, n_tests, ave_coeffs, pvals_coeffs, confidence_intervals
):
    """Generate and print a table."""
    model_rows = [header]

    for model in models:
        row = [OUTPUT_TABLES_MODEL_NAMES[model]]
        for i in range(5):
            coeff_value = ave_coeffs[model][i]
            coeff_symbol = significance_symbol(pvals_coeffs[model][i], n_tests=n_tests) if model not in [Recommender_random.__name__] else ""
            coeff_with_symbol = round_str(coeff_value) + coeff_symbol
            coeff_with_symbol_formatted = (
                (
                    "<bold>" if np.abs(coeff_value) == np.min(np.abs([ave_coeffs[m][i] for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]])) else 
                    "<underline>" if np.abs(coeff_value) in np.sort(np.abs([ave_coeffs[m][i] for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]]))[:2] else 
                    ""
                ) 
                + coeff_with_symbol
            )
            row.append(coeff_with_symbol_formatted)

        mean_coeff = np.mean(ave_coeffs[model])
        mean_coeff_formatted = (
            (
                "<bold>" if np.abs(mean_coeff) == np.min(np.abs([np.mean(ave_coeffs[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]])) else 
                "<underline>" if np.abs(mean_coeff) in np.sort(np.abs([np.mean(ave_coeffs[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]]))[:2] else 
                ""
            ) 
            + round_str(mean_coeff)
        )
        row.append(mean_coeff_formatted)
        model_rows.append(row)

        pval_row = [""]
        pval_row.extend([round_str(pvals_coeffs[model][i], 4) for i in range(5)])
        pval_row.append(" ")
        if model not in [Recommender_random.__name__,]:
            model_rows.append(pval_row)

        interval_row = [""]
        interval_row.extend(
            [
                "[" + round_str(interval[0]) + "," + round_str(interval[1]) + "]"
                for interval in confidence_intervals[model]
            ]
        )
        interval_row.append(" ")
        model_rows.append(interval_row)

    tab = np.asarray(model_rows)
    print(title)
    print_table(tab)

    return tab

def generate_and_print_performance_table(
    title,
    header,
    models,
    n_tests,
    NDCG_data_unbiased,
    NDCG_data_biased,
    pval_NDCGs_unbiased,
    pval_NDCGs_biased,
):
    """Generate and print a table for model performance."""
    model_rows = []

    for model in models:
        NDCG_unbiased_mean = np.mean(NDCG_data_unbiased[model])
        NDCG_unbiased_mean_with_symbol_formatted = (
            (
                "<bold>" if NDCG_unbiased_mean == np.max([np.mean(NDCG_data_unbiased[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]]) else
                "<underline>" if NDCG_unbiased_mean in np.sort([np.mean(NDCG_data_unbiased[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]])[-2:] else
                ""
            )
            + round_str(NDCG_unbiased_mean, 3)
            + (
                significance_symbol(pval_NDCGs_unbiased[model], n_tests=n_tests) if (
                    model not in [Recommender_random.__name__, Recommender_most_popular.__name__]
                    and NDCG_unbiased_mean > np.mean(NDCG_data_unbiased[Recommender_most_popular.__name__])
                ) else ""
            )
        )

        NDCG_biased_mean = np.mean(NDCG_data_biased[model])
        NDCG_biased_mean_with_symbol_formatted = (
            (
                "<bold>" if NDCG_biased_mean == np.max([np.mean(NDCG_data_biased[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]]) else
                "<underline>" if NDCG_biased_mean in np.sort([np.mean(NDCG_data_biased[m]) for m in models if m not in [Recommender_random.__name__, Recommender_most_popular.__name__]])[-2:] else
                ""
            )
            + round_str(NDCG_biased_mean, 3)
            + (
                significance_symbol(pval_NDCGs_biased[model], n_tests=n_tests) if (
                    model not in [Recommender_random.__name__, Recommender_most_popular.__name__]
                    and NDCG_biased_mean > np.mean(NDCG_data_biased[Recommender_most_popular.__name__])
                ) else ""
            )
        )

        NDCG_unbiased_ci = stats.bootstrap(
            (np.reshape(NDCG_data_unbiased[model], -1),),
            np.mean,
            confidence_level=0.99,
            vectorized=True,
        ).confidence_interval
        NDCG_biased_ci = stats.bootstrap(
            (np.reshape(NDCG_data_biased[model], -1),),
            np.mean,
            confidence_level=0.99,
            vectorized=True,
        ).confidence_interval

        model_rows.append(
            [
                OUTPUT_TABLES_MODEL_NAMES[model],
                NDCG_unbiased_mean_with_symbol_formatted,
                NDCG_biased_mean_with_symbol_formatted,
            ]
        )

        if model not in [Recommender_random.__name__, Recommender_most_popular.__name__]:
            model_rows.append(
                [""]
                + [round_str(pval_NDCGs_unbiased[model], 4)]
                + [round_str(pval_NDCGs_biased[model], 4)]
            )
        model_rows.append(
            [""]
            + [
                "["
                + round_str(NDCG_unbiased_ci[0], 3)
                + ","
                + round_str(NDCG_unbiased_ci[1], 3)
                + "]"
            ]
            + [
                "["
                + round_str(NDCG_biased_ci[0], 3)
                + ","
                + round_str(NDCG_biased_ci[1], 3)
                + "]"
            ]
        )
    rows = [header] + model_rows

    print(title)
    tab = np.asarray(rows)
    print_table(tab)

    return tab

def calculate_popularity_metrics(users, options, choices, all_set_types, set_type):
    """Calculate popularity metrics."""
    options_tmp = np.asarray(
        [options[i] for i in range(len(users)) if all_set_types[i] in [set_type]]
    )
    choices_tmp = np.asarray(
        [choices[i] for i in range(len(users)) if all_set_types[i] in [set_type]]
    )

    pickfreqs = np.unique(choices_tmp, return_counts=True)[1]
    expfreqs = np.unique(np.reshape(options_tmp, -1), return_counts=True)[1]

    rel = pickfreqs / expfreqs
    most_to_least_pop = [x for _, x in sorted(zip(rel, range(50)), reverse=True)]
    ranks = [most_to_least_pop.index(i) for i in range(5)]

    return rel, ranks

def generate_and_print_popularity_table(
    title, header
):
    """Generate and print a table for comparing popularity metrics."""
    rows = [header]

    for i in range(5):
        rows.append([i])

    (
        users,
        _,
        choices,
        options,
        all_set_types,
        _,
        _,
        _,
        _,
        _,
        _,
        _
    ) = load_dataset()

    for set_type in ["B", "BIAS"]:
        rel_pop, ranks_pop = calculate_popularity_metrics(
            users, options, choices, all_set_types, set_type
        )

        for i in range(5):
            rows[i + 1].extend([round_str(rel_pop[i]), round_str(ranks_pop[i])])

    for i in range(1, 6):
        rows[i].append(round_str(float(rows[i][1]) - float(rows[i][3])))
        rows[i].append(round_str(int(rows[i][2]) - int(rows[i][4])))

    print(title)
    tab = np.asarray(rows)
    print_table(tab)

    return tab

def saveTables(tables, save_path):
    """Save tables to a docx file."""
    print("Saving tables.")
    
    document = docx.Document()

    for tab in tables:
        nrows = len(tab)
        ncols = len(tab[0])

        table = document.add_table(rows=nrows, cols=ncols)

        content = tab[1:]
        cell_num = 0
        for idx, r in enumerate(content):
            # mergerows = int(
            #     1
            #     + len([entry for entry in content[:, 0] if entry == ""])
            #     / len([entry for entry in content[:, 0] if entry != ""])
            # )
            if content[idx][0] != "":
                mergerows = np.where(content[idx:, 0] != "")[0][1] if len(np.where(content[idx:, 0] != "")[0]) > 1 else len(content[idx:, 0])
                entries = [
                    "\n".join(content[(idx) : (idx + mergerows), i])
                    for i in range(len(r))
                ]
                for col in range(len(r)):
                    cell = table.cell(cell_num, col)
                    # table.cell(int(idx / mergerows), col).text = entries[col]
                    text_first_row = entries[col].split('\n')[0]
                    bold = False
                    underline = False
                    if len(text_first_row.split('<bold>')) > 1:
                        text_first_row = text_first_row.split('<bold>')[1]
                        bold = True
                    if len(entries[col].split('<underline>')) > 1:
                        text_first_row = text_first_row.split('<underline>')[1]
                        underline = True

                    paragraph = cell.paragraphs[0]

                    sentence = paragraph.add_run(text_first_row + "\n")
                    sentence.font.name = "Linux Libertine Display G"
                    sentence.font.size = docx.shared.Pt(8)
                    sentence.bold = bold
                    sentence.underline = underline

                    for line_num, line in enumerate(entries[col].split('\n')[1:]):
                        if len(entries[col].split('\n')) > 1:
                            sentence = paragraph.add_run(line + "\n" if line_num != len(entries[col].split('\n')[1:]) - 1 else line)
                            sentence.font.size = docx.shared.Pt(7) 
                            sentence.font.bold = False
                            sentence.font.underline = False
                    
                    cell.vertical_alignment = 1 # center
                    if col > 0:
                            paragraph.alignment = 1 # center
                cell_num += 1

    create_path(save_path)
    document.save(save_path)

def create_tables():
    """Create tables."""
    print("Creating tables...")

    with open(EVALUATION_RESULTS_PATH) as f:
        processed_results = json.load(f)

    raw_results = read_data_for_models()

    tables = []
    header = ["Model", "Item 1", "Item 2", "Item 3", "Item 4", "Item 5", "L2-Norm"]
    performance_header = ["Model", "nDCG Unbiased", "nDCG Biased"]

    models = [model for model in OUTPUT_TABLES_MODEL_NAMES if model in processed_results["bias"]["overexposure"].keys()]

    n_tests = (len(models) - 2) * 5
    performance_n_tests = (len(models) - 2) * 2

    ##### OVEREXPOSURE TABLES ########
    table = generate_and_print_bias_table(
        title="Bias Overexposure",
        header=header,
        n_tests=n_tests,
        models=models,
        ave_coeffs={model: processed_results["bias"]["overexposure"][model]["ave_bias_coeffs_per_item"] for model in models},
        pvals_coeffs={model: processed_results["bias"]["overexposure"][model]["pvals_bias_coeffs_per_item"] for model in models},
        confidence_intervals={model: processed_results["bias"]["overexposure"][model]["bias_confidence_intervals"] for model in models},
    )
    tables.append(table)

    table = generate_and_print_bias_table(
        title="Bias Overexposure Adjusted",
        header=header,
        models=models,
        n_tests=n_tests,
        ave_coeffs={model: processed_results["bias_adjusted"][model]["ave_adj_bias_coeffs"] for model in models},
        pvals_coeffs={model: processed_results["bias_adjusted"][model]["pvals_adj_bias_coeffs_per_item"] for model in models},
        confidence_intervals={model: processed_results["bias_adjusted"][model]["adj_bias_confidence_intervals"] for model in models},
    )
    tables.append(table)

    table = generate_and_print_performance_table(
        title="Performance Overexposure",
        header=performance_header,
        models=models,
        n_tests=performance_n_tests,
        NDCG_data_unbiased={model: raw_results["overexposure"][model]["mean_nDCG_B"] for model in models},
        NDCG_data_biased={model: raw_results["overexposure"][model]["mean_nDCG_BIAS"] for model in models},
        pval_NDCGs_unbiased={model: processed_results["performance"]["overexposure"][model]["pval_nDCG_B"] for model in models},
        pval_NDCGs_biased={model: processed_results["performance"]["overexposure"][model]["pval_nDCG_BIAS"] for model in models},
    )
    tables.append(table)


    ##### UNFAIR ALTERNATIVES TABLES ########
    table = generate_and_print_bias_table(
        title="Bias Unfair Alternatives",
        header=header,
        n_tests=n_tests,
        models=models,
        ave_coeffs={model: processed_results["bias"]["competition"][model]["ave_bias_coeffs_per_item"] for model in models},
        pvals_coeffs={model: processed_results["bias"]["competition"][model]["pvals_bias_coeffs_per_item"] for model in models},
        confidence_intervals={model: processed_results["bias"]["competition"][model]["bias_confidence_intervals"] for model in models},
    )
    tables.append(table)

    table = generate_and_print_performance_table(
        title="Performance Unfair Alternatives",
        header=performance_header,
        models=models,
        n_tests=performance_n_tests,
        NDCG_data_unbiased={model: raw_results["competition"][model]["mean_nDCG_popular"] for model in models},
        NDCG_data_biased={model: raw_results["competition"][model]["mean_nDCG_unpopular"] for model in models},
        pval_NDCGs_unbiased={model: processed_results["performance"]["competition"][model]["pval_nDCG_popular"] for model in models},
        pval_NDCGs_biased={model: processed_results["performance"]["competition"][model]["pval_nDCG_unpopular"] for model in models},
    )
    tables.append(table)

    ######## POPULARITY TABLE OF ITEMS FROM B AND BIAS ########
    header = ["Course ID", "Choice ratio B", "Choice rank B", 
              "Choice ratio Bias", "Choice rank Bias", 
              "Choice ratio Diff", "Choice rank Diff"]
    table = generate_and_print_popularity_table(
        title="Popularity Comparison", header=header
    )
    tables.append(table)

    saveTables(tables, OUTPUT_TABLE_PATH)